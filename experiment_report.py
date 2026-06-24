from __future__ import annotations

import csv
import os
import shutil
from datetime import datetime


def _safe_name(x):
    x = str(x).strip().lower()
    x = x.replace(" ", "_")
    x = x.replace("-", "_")
    x = x.replace("/", "_")
    return x


def _get_hr_mode(config):
    candidates = [
        "HR_MODE",
        "HR_REGIME",
        "HR_DYNAMICS_MODE",
        "HINDMARSH_ROSE_MODE",
    ]

    for name in candidates:
        if hasattr(config, name):
            value = getattr(config, name)
            if value:
                return _safe_name(value)

    return "hr_run"


def _clear_folder(folder):
    if not os.path.exists(folder):
        return

    for name in os.listdir(folder):
        path = os.path.join(folder, name)

        if os.path.isfile(path) or os.path.islink(path):
            os.remove(path)

        elif os.path.isdir(path):
            shutil.rmtree(path)


def setup_run_output_folder(config):
    """
    Creates one clean output folder per HR regime.

    Example:
        outputs/periodic_spiking/
        outputs/periodic_bursting/
        outputs/chaotic_bursting/
    """

    root_output = getattr(config, "OUTPUT_ROOT", "outputs")
    os.makedirs(root_output, exist_ok=True)

    dataset_mode = getattr(config, "DATASET_MODE", "hr")

    if dataset_mode == "hr":
        run_name = _get_hr_mode(config)
    else:
        run_name = _safe_name(dataset_mode)

    run_output_dir = os.path.join(root_output, run_name)

    clear_each_run = getattr(config, "CLEAR_OUTPUT_FOLDER_EACH_RUN", False)

    if clear_each_run:
        _clear_folder(run_output_dir)

    os.makedirs(run_output_dir, exist_ok=True)

    config.OUTPUT_DIR = run_output_dir

    return run_output_dir


def _preferred_column_order():
    return [
        "timestamp",
        "regime",
        "optimizer",
        "samples",
        "states",
        "rmse_x",
        "nrmse_x",
        "rmse_all_states",
        "nrmse_all_states",
        "N_res",
        "density_p",
        "spectral_radius",
        "leak_rate",
        "input_scaling",
        "ridge",
        "washout",
    ]


def _make_fieldnames(rows):
    """
    Important fix:
    Existing CSV files may have old columns.
    New rows may have new columns.
    So we build a union of all columns instead of using only row[0].keys().
    """

    preferred = _preferred_column_order()
    all_keys = set()

    for row in rows:
        all_keys.update(row.keys())

    final_keys = []

    for key in preferred:
        if key in all_keys:
            final_keys.append(key)

    for key in sorted(all_keys):
        if key not in final_keys:
            final_keys.append(key)

    return final_keys


def _write_csv(path, rows):
    if not rows:
        return

    os.makedirs(os.path.dirname(path), exist_ok=True)

    fieldnames = _make_fieldnames(rows)

    with open(path, "w", newline="") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=fieldnames,
            extrasaction="ignore",
        )
        writer.writeheader()

        for row in rows:
            clean_row = {key: row.get(key, "") for key in fieldnames}
            writer.writerow(clean_row)


def _read_csv(path):
    if not os.path.exists(path):
        return []

    try:
        with open(path, newline="") as f:
            return list(csv.DictReader(f))
    except Exception:
        return []


def _upsert_global_row(path, row, key="regime"):
    rows = _read_csv(path)

    found = False
    new_rows = []

    for old in rows:
        if old.get(key) == row.get(key):
            new_rows.append(row)
            found = True
        else:
            new_rows.append(old)

    if not found:
        new_rows.append(row)

    _write_csv(path, new_rows)


def _try_save_docx(path, rows, title):
    try:
        from docx import Document
    except Exception:
        print("[Report] python-docx not installed. DOCX report skipped.")
        print("[Report] Install with: pip install python-docx")
        return False

    if not rows:
        return False

    os.makedirs(os.path.dirname(path), exist_ok=True)

    fieldnames = _make_fieldnames(rows)

    doc = Document()
    doc.add_heading(title, level=1)

    table = doc.add_table(rows=1, cols=len(fieldnames))
    table.style = "Table Grid"

    header_cells = table.rows[0].cells
    for i, key in enumerate(fieldnames):
        header_cells[i].text = str(key)

    for row in rows:
        cells = table.add_row().cells
        for i, key in enumerate(fieldnames):
            cells[i].text = str(row.get(key, ""))

    doc.save(path)
    return True


def _save_optimizer_table(base_output_dir, optimizer_summary):
    if not optimizer_summary:
        return

    rows = []

    for r in optimizer_summary:
        rows.append(
            {
                "rank": len(rows) + 1,
                "optimizer": r.get("optimizer", ""),
                "best_score": round(float(r.get("best_score", 0.0)), 8),
                "validation_nrmse_x": round(float(r.get("validation_nrmse_x", 0.0)), 8),
                "validation_nrmse": round(float(r.get("validation_nrmse", 0.0)), 8),
                "N_res": int(r.get("N_res", 0)),
                "density_p": round(float(r.get("p", 0.0)), 6),
                "spectral_radius": round(float(r.get("spectral_radius", 0.0)), 6),
                "leak_rate": round(float(r.get("leaky_coefficient", 0.0)), 6),
                "input_scaling": round(float(r.get("input_scaling", 0.0)), 6),
                "ridge": float(r.get("regularization", 0.0)),
                "washout": int(r.get("washout", 0)),
            }
        )

    csv_path = os.path.join(base_output_dir, "optimizer_ranking_table.csv")
    docx_path = os.path.join(base_output_dir, "optimizer_ranking_table.docx")

    _write_csv(csv_path, rows)
    _try_save_docx(docx_path, rows, "Optimizer Ranking Table")

    print(f"[Report] Saved optimizer ranking table -> {csv_path}")


def save_experiment_summary(
    loader,
    config,
    optimizer_name,
    params,
    metrics,
    base_output_dir,
    optimizer_summary=None,
):
    """
    Saves clean report files.

    Inside each regime folder:
        experiment_summary.csv
        experiment_summary.docx
        optimizer_ranking_table.csv
        optimizer_ranking_table.docx

    Inside outputs root:
        experiment_comparison.csv
        experiment_comparison.docx
    """

    os.makedirs(base_output_dir, exist_ok=True)

    root_output = os.path.dirname(base_output_dir)
    regime = os.path.basename(base_output_dir)

    row = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "regime": regime,
        "optimizer": optimizer_name,
        "samples": getattr(loader, "n_samples", ""),
        "states": getattr(loader, "n_neurons", ""),
        "rmse_x": round(float(metrics.get("rmse_x", 0.0)), 8),
        "nrmse_x": round(float(metrics.get("nrmse_x", 0.0)), 8),
        "rmse_all_states": round(float(metrics.get("rmse_all_states", 0.0)), 8),
        "nrmse_all_states": round(float(metrics.get("nrmse_all_states", 0.0)), 8),
        "N_res": int(params.get("N_res", 0)),
        "density_p": round(float(params.get("p", 0.0)), 6),
        "spectral_radius": round(float(params.get("spectral_radius", 0.0)), 6),
        "leak_rate": round(float(params.get("leaky_coefficient", 0.0)), 6),
        "input_scaling": round(float(params.get("input_scaling", 0.0)), 6),
        "ridge": float(params.get("regularization", 0.0)),
        "washout": int(params.get("washout", 0)),
    }

    single_csv = os.path.join(base_output_dir, "experiment_summary.csv")
    single_docx = os.path.join(base_output_dir, "experiment_summary.docx")

    _write_csv(single_csv, [row])
    _try_save_docx(single_docx, [row], "Single Experiment Summary")

    global_csv = os.path.join(root_output, "experiment_comparison.csv")
    global_docx = os.path.join(root_output, "experiment_comparison.docx")

    _upsert_global_row(global_csv, row, key="regime")

    global_rows = _read_csv(global_csv)
    _try_save_docx(global_docx, global_rows, "Experiment Comparison Across HR Regimes")

    _save_optimizer_table(base_output_dir, optimizer_summary or [])

    print(f"[Report] Saved single-run summary -> {single_csv}")
    print(f"[Report] Saved global comparison -> {global_csv}")

# ============================================================
# CONTROL REPORT HELPERS
# Append this block near the bottom of experiment_report.py
# ============================================================

def _control_fieldnames(rows):
    preferred = [
        "timestamp",
        "regime",
        "controller",
        "best_K",
        "best_target_rmse_state",
        "best_target_rmse_x",
        "best_settling_time_s",
        "best_spike_reduction_percent",
        "best_control_energy_mean",
        "best_diverged",
        "target_mode",
    ]

    all_keys = set()
    for row in rows:
        all_keys.update(row.keys())

    fieldnames = []
    seen = set()

    for key in preferred:
        if key in all_keys:
            fieldnames.append(key)
            seen.add(key)

    for key in sorted(all_keys):
        if key not in seen:
            fieldnames.append(key)

    return fieldnames


def _control_write_csv(path, rows):
    if not rows:
        return

    os.makedirs(os.path.dirname(path), exist_ok=True)
    fieldnames = _control_fieldnames(rows)

    with open(path, "w", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        for row in rows:
            writer.writerow({k: row.get(k, "") for k in fieldnames})


def _control_read_csv(path):
    if not os.path.exists(path):
        return []

    try:
        with open(path, newline="") as f:
            return list(csv.DictReader(f))
    except Exception:
        return []


def _control_try_docx(path, rows, title):
    try:
        from docx import Document
    except Exception:
        print("[Control Report] python-docx not installed. DOCX export skipped.")
        return False

    if not rows:
        return False

    os.makedirs(os.path.dirname(path), exist_ok=True)
    fieldnames = _control_fieldnames(rows)

    doc = Document()
    doc.add_heading(title, level=1)

    table = doc.add_table(rows=1, cols=len(fieldnames))
    table.style = "Table Grid"

    for i, key in enumerate(fieldnames):
        table.rows[0].cells[i].text = str(key)

    for row in rows:
        cells = table.add_row().cells
        for i, key in enumerate(fieldnames):
            cells[i].text = str(row.get(key, ""))

    doc.save(path)
    return True


def save_control_summary(regime_output_dir, controller_name, sweep_rows):
    """
    Save:
      1. per-regime control summary inside outputs/<regime>/control/<controller_name>/
      2. global control comparison inside outputs/control_comparison.csv and .docx
    """
    if not sweep_rows:
        return

    regime_output_dir = os.path.abspath(regime_output_dir)
    regime_name = os.path.basename(regime_output_dir)
    root_output = os.path.dirname(regime_output_dir)
    controller_dir = os.path.join(regime_output_dir, "control", controller_name)

    os.makedirs(controller_dir, exist_ok=True)

    valid = [r for r in sweep_rows if not bool(r.get("diverged", False))]
    candidates = valid if valid else sweep_rows

    def _score(row):
        try:
            return (
                float(row.get("target_rmse_state", 1e18)),
                float(row.get("control_energy_mean", 1e18)),
                -float(row.get("spike_reduction_percent", -1e18)),
            )
        except Exception:
            return (1e18, 1e18, 1e18)

    best = min(candidates, key=_score)

    best_row = {
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "regime": regime_name,
        "controller": controller_name,
        "best_K": best.get("K", ""),
        "best_target_rmse_state": best.get("target_rmse_state", ""),
        "best_target_rmse_x": best.get("target_rmse_x", ""),
        "best_settling_time_s": best.get("settling_time_s", ""),
        "best_spike_reduction_percent": best.get("spike_reduction_percent", ""),
        "best_control_energy_mean": best.get("control_energy_mean", ""),
        "best_diverged": best.get("diverged", ""),
        "target_mode": best.get("target_mode", ""),
    }

    # Per-regime summary
    single_csv = os.path.join(controller_dir, "control_experiment_summary.csv")
    single_docx = os.path.join(controller_dir, "control_experiment_summary.docx")

    _control_write_csv(single_csv, [best_row])
    _control_try_docx(single_docx, [best_row], "Control Experiment Summary")

    # Global comparison
    global_csv = os.path.join(root_output, "control_comparison.csv")
    global_docx = os.path.join(root_output, "control_comparison.docx")

    existing = _control_read_csv(global_csv)
    new_rows = []
    updated = False

    for old in existing:
        if (
            str(old.get("regime", "")) == str(best_row["regime"])
            and str(old.get("controller", "")) == str(best_row["controller"])
        ):
            new_rows.append(best_row)
            updated = True
        else:
            new_rows.append(old)

    if not updated:
        new_rows.append(best_row)

    _control_write_csv(global_csv, new_rows)
    _control_try_docx(global_docx, new_rows, "Control Comparison Across Regimes")

    print(f"[Control Report] Saved per-regime summary -> {single_csv}")
    print(f"[Control Report] Saved global comparison -> {global_csv}")
